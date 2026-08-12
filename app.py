import streamlit as st
import pandas as pd
from datetime import datetime
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

st.set_page_config(page_title="Inventario DDUMA - Gestión de Bienes y Oficios", layout="wide")

st.title("📋 Control de Inventario y Generador de Oficios - DDUMA")
st.write("Sube tu archivo Excel de inventario para verificar, reasignar resguardantes y generar oficios formales.")

# 1. Cargar archivo Excel
uploaded_file = st.file_uploader("📂 Sube aquí tu archivo Excel de Inventario", type=["xlsx", "xls"])

if uploaded_file is not None:
    try:
        # Cargar datos en estado de sesión para mantener los cambios de resguardante
        if "df_inventario" not in st.session_state or st.sidebar.button("🔄 Recargar Excel Original"):
            df_raw = pd.read_excel(uploaded_file)
            df_raw.columns = df_raw.columns.astype(str).str.strip()
            st.session_state.df_inventario = df_raw.copy()

        df_current = st.session_state.df_inventario

        # Identificar la columna del resguardante
        col_resguardante = None
        for col in df_current.columns:
            if "resguardante actual" in col.lower() or "resguardante_actual" in col.lower():
                col_resguardante = col
                break
        
        if col_resguardante is None:
            for col in df_current.columns:
                if "actual" in col.lower() or "resguardante" in col.lower():
                    col_resguardante = col
                    break

        if col_resguardante is not None:
            # Configuración general de la Dirección Destinataria
            st.subheader("⚙️ Configuración del Destinatario y Tipo de Oficio")
            
            col_cfg1, col_cfg2, col_cfg3 = st.columns(3)

            with col_cfg1:
                nombre_direccion_dest = st.text_input("Dirección Destinataria", value="CATASTRO")
                limpio_dest = nombre_direccion_dest.upper().replace("DIRECCIÓN DE", "").replace("DIRECCION DE", "").strip()
                area_asignada_texto = f"DIRECCION DE {limpio_dest}"

            with col_cfg2:
                destinatario_nombre = st.text_input("Director / Titular Destinatario", value="LIC. JOSÉ DOMINGO [APELLIDOS]")
                num_oficio = st.text_input("Número de Oficio", value="0542/DDUMA/2026")

            with col_cfg3:
                remitente_nombre = st.text_input("Remitente (Titular DDUMA)", value="LIC. ROSENDO SÁNCHEZ PREVE")
                remitente_cargo = st.text_input("Cargo Remitente", value="DIRECTORA DE DESARROLLO URBANO Y MEDIO AMBIENTE")
                fecha_oficio = st.date_input("Fecha del Oficio", value=datetime.now())

            st.markdown("---")

            # Filtrar datos de la área seleccionada
            resguardante_clean = df_current[col_resguardante].astype(str).str.upper().str.strip()
            filtro = resguardante_clean.str.contains(limpio_dest, na=False)
            df_filtrado = df_current[filtro].copy()

            if len(df_filtrado) > 0:
                st.success(f"✅ Se encontraron {len(df_filtrado)} bienes relacionados con {area_asignada_texto}.")

                # Pestanas interactiva: Reasignar vs Oficio
                tab1, tab2 = st.tabs(["👤 Reasignar Compañero / Resguardante", "📄 Generar Oficio de Verificación"])

                # TAB 1: REASIGNACIÓN DIRECTA
                with tab1:
                    st.write("### Asignar un compañero o resguardante específico a los bienes")
                    
                    col_mod1, col_mod2 = st.columns([2, 1])
                    with col_mod1:
                        # Identificar columna clave (Clave Inventario o primera columna)
                        col_clave = df_filtrado.columns[0]
                        bienes_opciones = df_filtrado[col_clave].astype(str) + " - " + df_filtrado.iloc[:, 1].astype(str)
                        bien_seleccionado = st.selectbox("Selecciona el bien a reasignar:", bienes_opciones)
                        clave_elegida = bien_seleccionado.split(" - ")[0]

                    with col_mod2:
                        nuevo_resguardante = st.text_input("Nombre del nuevo resguardante (Ej. Juan Pérez - Catastro)")
                        if st.button("✏️ Actualizar Resguardante"):
                            if nuevo_resguardante.strip():
                                idx_target = st.session_state.df_inventario[
                                    st.session_state.df_inventario[col_clave].astype(str) == clave_elegida
                                ].index
                                st.session_state.df_inventario.loc[idx_target, col_resguardante] = nuevo_resguardante.strip().upper()
                                st.success(f"¡Resguardante actualizado a '{nuevo_resguardante.upper()}'!")
                                st.rerun()

                    st.write("#### Lista Actualizada de Bienes")
                    df_bienes_editados = st.data_editor(df_filtrado, use_container_width=True, key="editor_bienes")

                # TAB 2: GENERACIÓN DEL OFICIO DE VERIFICACIÓN
                with tab2:
                    st.write("### Solicitud de Verificación Física de Bienes")
                    
                    tipo_oficio = st.radio(
                        "Selecciona el motivo del oficio:",
                        [
                            "1. Solicitud de Verificación Física en Campo (Confirmar si los bienes se encuentran físicamente en su área)",
                            "2. Solicitud de Regularización y Cambio Formal de Resguardo"
                        ]
                    )

                    # Función para Generar el Word
                    def generar_word_verificacion(df_data):
                        doc = docx.Document()

                        for section in doc.sections:
                            section.top_margin = Inches(0.8)
                            section.bottom_margin = Inches(0.8)
                            section.left_margin = Inches(0.8)
                            section.right_margin = Inches(0.8)

                        # Lema
                        p_lema = doc.add_paragraph('"2026, Año de Margarita Maza Parada"')
                        p_lema.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_lema.runs[0].font.italic = True
                        p_lema.runs[0].font.size = Pt(8.5)

                        # Encabezado
                        p_meta = doc.add_paragraph()
                        p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        p_meta.add_run("DIRECCIÓN DE DESARROLLO URBANO Y MEDIO AMBIENTE\n").bold = True
                        p_meta.add_run(f"OFICIO: {num_oficio}\n").bold = True
                        
                        asunto_txt = "Solicitud de verificación e inspección física de bienes muebles en sus instalaciones." if "1" in tipo_oficio else f"Notificación y solicitud de regularización de resguardos en la {area_asignada_texto}."
                        p_meta.add_run(f"ASUNTO: {asunto_txt}\n").bold = True

                        # Fecha
                        meses = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"]
                        fecha_fmt = f"{fecha_oficio.day} de {meses[fecha_oficio.month - 1]} del {fecha_oficio.year}"
                        p_fecha = doc.add_paragraph(f"San Francisco de Campeche, Camp., a {fecha_fmt}\n")

                        # Destinatario
                        p_dest = doc.add_paragraph()
                        p_dest.add_run(f"{destinatario_nombre}\n").bold = True
                        p_dest.add_run(f"DIRECTOR DE {limpio_dest}\n").bold = True
                        p_dest.add_run("P R E S E N T E .-")
                        p_dest.paragraph_format.space_after = Pt(12)

                        # Cuerpo dinámico según la opción elegida
                        if "1" in tipo_oficio:
                            parrafos = [
                                "Me dirijo a usted de la manera más atenta en el marco de las actividades de control, seguimiento y actualización del inventario patrimonial de este H. Ayuntamiento.",
                                f"Sobre el particular, le solicito atentamente su valioso apoyo a efecto de realizar la verificación y constatación física en campo de los bienes muebles que a continuación se relacionan, los cuales se tienen ubicados preliminarmente en las instalaciones y áreas operativas de la {area_asignada_texto} a su digno cargo.",
                                "Agradeceré se sirva confirmar si dichos bienes se encuentran efectivamente en su área y, en su caso, nos proporcione el nombre del servidor público que los tiene bajo su resguardo u operatividad directa, a fin de proceder con la actualización de los registros institucionales.",
                                "La relación de los bienes a verificar se detalla a continuación:"
                            ]
                        else:
                            parrafos = [
                                "Me dirijo a usted de la manera más atenta con relación a las acciones de verificación y depuración del inventario físico de esta Dirección de Desarrollo Urbano y Medio Ambiente.",
                                f"Hago de su conocimiento que se ha identificado la presencia de diversos bienes muebles patrimoniales asignados a la {area_asignada_texto} a su digno cargo.",
                                "En virtud de lo anterior, le solicito atentamente girar sus apreciables instrucciones para proceder con la formalización del cambio de resguardo en el Catálogo General de Bienes Muebles del Municipio.",
                                "A continuación, se detalla la relación de los bienes muebles antes referidos:"
                            ]

                        for p_text in parrafos:
                            p = doc.add_paragraph(p_text)
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.space_after = Pt(6)
                            p.paragraph_format.line_spacing = 1.15

                        # Tabla de Bienes en Word
                        table = doc.add_table(rows=1, cols=len(df_data.columns))
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        hdr_cells = table.rows[0].cells
                        for col_idx, col_name in enumerate(df_data.columns):
                            hdr_cells[col_idx].text = str(col_name).upper()
                            hdr_cells[col_idx].paragraphs[0].runs[0].font.bold = True
                            hdr_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(8)

                        for idx, row in df_data.iterrows():
                            row_cells = table.add_row().cells
                            for col_idx, col_name in enumerate(df_data.columns):
                                val = str(row[col_name]) if pd.notna(row[col_name]) else ""
                                row_cells[col_idx].text = val
                                row_cells[col_idx].paragraphs[0].runs[0].font.size = Pt(8)

                        # Firma
                        p_cierre = doc.add_paragraph("\nSin otro particular por el momento, le reitero la seguridad de mi atenta y distinguida consideración.\n")
                        p_cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        p_atentamente = doc.add_paragraph("A T E N T A M E N T E\n\n\n\n")
                        p_atentamente.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_atentamente.add_run(f"{remitente_nombre}\n").bold = True
                        p_atentamente.add_run(f"{remitente_cargo}").bold = True

                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        return buffer

                    # Botón de Descarga del Oficio
                    st.download_button(
                        label="📄 Descargar Oficio Personalizado (.docx)",
                        data=generar_word_verificacion(df_filtrado),
                        file_name=f"Oficio_Verificacion_{limpio_dest}_{num_oficio.replace('/', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning(f"⚠️ No se encontraron bienes relacionados con '{limpio_dest}'.")
        else:
            st.error("❌ No se encontró la columna de resguardante en el archivo Excel.")
            
    except Exception as e:
        st.error(f"Error al procesar el archivo Excel: {e}")
else:
    st.info("👆 Por favor sube tu archivo Excel de inventario para comenzar.")
